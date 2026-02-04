"""
Скрипт для извлечения текста из DOCX Приказа №281/пр
"""

from docx import Document
from pathlib import Path

def extract_docx():
    docx_path = Path(__file__).parent / "prikaz_minstroya_rossii_ot_12.05.2025_n_281_pr_o_normativnyh.docx"
    
    if not docx_path.exists():
        print(f"Файл не найден: {docx_path}")
        return
    
    doc = Document(str(docx_path))
    
    # Извлекаем параграфы
    text_content = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)
    
    # Извлекаем таблицы
    table_content = []
    for i, table in enumerate(doc.tables):
        table_content.append(f"\n=== ТАБЛИЦА {i+1} ===")
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.replace("|", "").strip():
                table_content.append(row_text)
    
    # Сохраняем
    output_path = Path(__file__).parent / "docx_extracted.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("=== ПАРАГРАФЫ ===\n\n")
        f.write("\n".join(text_content))
        f.write("\n\n=== ТАБЛИЦЫ ===\n")
        f.write("\n".join(table_content))
    
    print(f"Извлечено {len(text_content)} параграфов и {len(doc.tables)} таблиц")
    print(f"Сохранено в: {output_path}")

if __name__ == "__main__":
    extract_docx()
