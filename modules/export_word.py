"""
Экспорт сметы в Word (DOCX)
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color: str):
    """Установить заливку ячейки"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def export_to_word(estimate, filename: str = None) -> Path:
    """Экспортировать смету в Word"""
    
    if filename is None:
        filename = f"Смета_{estimate.project_name}_{estimate.date_created}.docx"
    
    output_path = Path(filename)
    
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Заголовок
    title = doc.add_heading('ЛОКАЛЬНАЯ СМЕТА', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('на инженерно-геологические изыскания')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Информация о проекте
    info_data = [
        ("Проект:", estimate.project_name),
        ("Шифр:", estimate.project_code or "-"),
        ("Объект:", estimate.object_name or "-"),
        ("Заказчик:", estimate.customer or "-"),
        ("Подрядчик:", estimate.contractor or "-"),
        ("Базовый город:", getattr(estimate, 'base_city', 'г. Санкт-Петербург')),
        ("Регион производства работ:", getattr(estimate, 'work_region', '-') or '-'),
        ("Расстояние до объекта:", f"{getattr(estimate, 'distance_km', '-')} км"),
        ("Дата:", estimate.date_created),
        ("Индекс пересчёта:", f"{float(estimate.price_index):.2f}"),
    ]
    
    info_table = doc.add_table(rows=len(info_data), cols=2)
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        info_table.rows[i].cells[1].text = str(value)
    
    info_table.columns[0].width = Cm(5)
    info_table.columns[1].width = Cm(11)
    
    doc.add_paragraph()
    
    # Основная таблица сметы — группируем по категориям
    categories = {
        "field": {"name": "ПОЛЕВЫЕ РАБОТЫ", "items": []},
        "laboratory": {"name": "ЛАБОРАТОРНЫЕ РАБОТЫ", "items": []},
        "office": {"name": "КАМЕРАЛЬНЫЕ РАБОТЫ", "items": []}
    }
    
    for item in estimate.items:
        if item.code.startswith(("01", "02", "03", "04")):
            categories["field"]["items"].append(item)
        elif item.code.startswith(("05", "06", "07")):
            categories["laboratory"]["items"].append(item)
        else:
            categories["office"]["items"].append(item)
    
    # Подсчитываем количество строк
    total_rows = 1  # Заголовок таблицы
    for cat_data in categories.values():
        if cat_data["items"]:
            total_rows += 1  # Заголовок раздела
            total_rows += len(cat_data["items"])
            total_rows += 1  # Подитог
    
    # ДЗ строки
    dz_costs = estimate.additional_costs or []
    total_rows += 1  # Итого базовые
    if dz_costs:
        total_rows += 1  # Заголовок "Дополнительные затраты"
        total_rows += len(dz_costs)
        total_rows += 1  # Итого с ДЗ
    
    total_rows += 1  # Итого с индексом
    if float(estimate.contract_coefficient) != 1.0:
        total_rows += 1  # Кдог
    total_rows += 1  # ВСЕГО
    
    table = doc.add_table(rows=total_rows, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ширина колонок
    widths = [Cm(1), Cm(1.5), Cm(7), Cm(1.5), Cm(1.5), Cm(2), Cm(2.5)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width
    
    # Заголовок таблицы
    headers = ["№", "Код", "Наименование работ", "Ед.", "Кол-во", "Цена", "Сумма"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9D9D9')
    
    current_row = 1
    item_num = 1
    
    for cat_key, cat_data in categories.items():
        if not cat_data["items"]:
            continue
        
        # Заголовок раздела
        row = table.rows[current_row]
        row.cells[0].merge(row.cells[6])
        row.cells[0].text = cat_data["name"]
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], 'E0E0E0')
        current_row += 1
        
        # Позиции
        for item in cat_data["items"]:
            row = table.rows[current_row]
            row.cells[0].text = str(item_num)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row.cells[1].text = item.code
            row.cells[2].text = item.name
            row.cells[3].text = item.unit
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row.cells[4].text = f"{float(item.quantity):.1f}"
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Для рекогносцировки — показываем формулу
            if float(item.pz1p_fixed) > 0:
                pz1p = float(item.pz1p_fixed)
                pz2p = float(item.base_cost)
                row.cells[5].text = f"ПЗ1п({pz1p:,.0f})+ПЗ2п({pz2p:,.0f})"
            else:
                row.cells[5].text = f"{float(item.unit_cost):,.0f}"
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            row.cells[6].text = f"{float(item.total_cost):,.0f}"
            row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            item_num += 1
            current_row += 1
        
        # Подитог раздела
        subtotal = sum(float(item.total_cost) for item in cat_data["items"])
        row = table.rows[current_row]
        row.cells[0].merge(row.cells[5])
        row.cells[0].text = f"Итого {cat_data['name'].lower()}:"
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[0].paragraphs[0].runs[0].bold = True
        
        row.cells[6].text = f"{subtotal:,.0f}"
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[6].paragraphs[0].runs[0].bold = True
        
        set_cell_shading(row.cells[0], 'FFF3E0')
        set_cell_shading(row.cells[6], 'FFF3E0')
        current_row += 1
    
    # Итого базовые
    base_total = sum(float(item.total_cost) for item in estimate.items)
    row = table.rows[current_row]
    row.cells[0].merge(row.cells[5])
    row.cells[0].text = "ИТОГО базовые затраты (СП + СЛ + СК):"
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[0].paragraphs[0].runs[0].bold = True
    
    row.cells[6].text = f"{base_total:,.0f}"
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[6].paragraphs[0].runs[0].bold = True
    
    set_cell_shading(row.cells[0], 'E8F5E9')
    set_cell_shading(row.cells[6], 'E8F5E9')
    current_row += 1
    
    # Дополнительные затраты
    dz_sum = 0
    if dz_costs:
        row = table.rows[current_row]
        row.cells[0].merge(row.cells[6])
        row.cells[0].text = "Дополнительные затраты:"
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].italic = True
        current_row += 1
        
        for dz_num, cost in enumerate(dz_costs, 1):
            row = table.rows[current_row]
            
            row.cells[0].text = f"ДЗ-{dz_num}"
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row.cells[1].merge(row.cells[2])
            row.cells[1].text = cost.get('name', 'ДЗ')
            
            row.cells[3].merge(row.cells[4])
            row.cells[3].text = cost.get('basis', '-')
            
            row.cells[5].text = cost.get('formula', '-')
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            val = round(float(cost.get('value', 0)), 2)
            dz_sum += val
            row.cells[6].text = f"{val:,.0f}"
            row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            current_row += 1
        
        # Итого с ДЗ
        total_with_dz = round(base_total + dz_sum, 2)
        row = table.rows[current_row]
        row.cells[0].merge(row.cells[5])
        row.cells[0].text = "ИТОГО с учётом дополнительных затрат:"
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[0].paragraphs[0].runs[0].bold = True
        
        row.cells[6].text = f"{total_with_dz:,.0f}"
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[6].paragraphs[0].runs[0].bold = True
        
        set_cell_shading(row.cells[0], 'FFF3E0')
        set_cell_shading(row.cells[6], 'FFF3E0')
        current_row += 1
    else:
        total_with_dz = base_total
    
    # С индексом пересчёта
    idx = float(estimate.price_index)
    total_indexed = round(total_with_dz * idx, 2)
    
    row = table.rows[current_row]
    row.cells[0].merge(row.cells[5])
    row.cells[0].text = f"ИТОГО с индексом пересчёта ({idx:.2f}):"
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[0].paragraphs[0].runs[0].bold = True
    
    row.cells[6].text = f"{total_indexed:,.0f}"
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[6].paragraphs[0].runs[0].bold = True
    
    set_cell_shading(row.cells[0], 'E8F5E9')
    set_cell_shading(row.cells[6], 'E8F5E9')
    current_row += 1
    
    # Коэффициент договорной цены
    k_contract = float(estimate.contract_coefficient)
    if k_contract != 1.0:
        final_total = round(total_indexed * k_contract, 2)
        
        row = table.rows[current_row]
        row.cells[0].merge(row.cells[5])
        row.cells[0].text = f"Коэффициент договорной цены ({k_contract:.3f}):"
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[0].paragraphs[0].runs[0].bold = True
        
        row.cells[6].text = f"{final_total:,.0f}"
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[6].paragraphs[0].runs[0].bold = True
        current_row += 1
    else:
        final_total = total_indexed
    
    # ВСЕГО ПО СМЕТЕ
    row = table.rows[current_row]
    row.cells[0].merge(row.cells[5])
    row.cells[0].text = "ВСЕГО по смете:"
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[0].paragraphs[0].runs[0].bold = True
    
    row.cells[6].text = f"{final_total:,.0f}"
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[6].paragraphs[0].runs[0].bold = True
    
    set_cell_shading(row.cells[0], 'C8E6C9')
    set_cell_shading(row.cells[6], 'C8E6C9')
    
    # Подписи
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("Составил: __________________ / __________________ /")
    doc.add_paragraph()
    doc.add_paragraph("Проверил: __________________ / __________________ /")
    
    # Сохранение
    doc.save(output_path)
    
    return output_path
